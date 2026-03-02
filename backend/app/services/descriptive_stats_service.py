import io
from typing import Any, Dict, List, Optional

import numpy as np
import pandas as pd
from scipy import stats as scipy_stats


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _get_histogram_bins(series: pd.Series, n_bins: int = 12) -> List[Dict[str, Any]]:
    clean = series.dropna()
    if len(clean) < 2:
        return []
    hist, edges = np.histogram(clean, bins=n_bins)
    return [
        {
            "bin_start": round(float(edges[i]), 4),
            "bin_end": round(float(edges[i + 1]), 4),
            "count": int(hist[i]),
        }
        for i in range(len(hist))
    ]


# ---------------------------------------------------------------------------
# Column-type detection
# ---------------------------------------------------------------------------

def detect_column_types(df: pd.DataFrame) -> List[Dict[str, Any]]:
    """Return metadata for every column: name, detected type, missing counts."""
    columns: List[Dict[str, Any]] = []
    for col in df.columns:
        n_total = len(df[col])
        n_missing = int(df[col].isna().sum())
        pct_missing = round(float(n_missing / n_total * 100), 1) if n_total > 0 else 0.0

        if pd.api.types.is_datetime64_any_dtype(df[col]):
            col_type = "date"
        elif pd.api.types.is_bool_dtype(df[col]):
            col_type = "categorical"
        elif pd.api.types.is_numeric_dtype(df[col]):
            n_unique = int(df[col].nunique())
            col_type = "continuous" if n_unique > 10 else "categorical"
        else:
            # Try to parse as numeric
            try:
                numeric_series = pd.to_numeric(df[col], errors="coerce")
                total_non_null = df[col].notna().sum()
                valid_count = numeric_series.notna().sum()
                if total_non_null > 0 and valid_count / total_non_null > 0.8:
                    n_unique = int(numeric_series.nunique())
                    col_type = "continuous" if n_unique > 10 else "categorical"
                else:
                    col_type = "categorical"
            except Exception:
                col_type = "categorical"

        columns.append(
            {
                "name": col,
                "type": col_type,
                "n_missing": n_missing,
                "pct_missing": pct_missing,
            }
        )
    return columns


# ---------------------------------------------------------------------------
# Per-variable statistics
# ---------------------------------------------------------------------------

def compute_variable_stats_categorical(
    df: pd.DataFrame, variable_name: str
) -> Dict[str, Any]:
    if variable_name not in df.columns:
        raise ValueError(f"Column '{variable_name}' not found in dataset")

    col = df[variable_name]
    n_total = int(len(col))
    n_missing = int(col.isna().sum())
    pct_missing = round(float(n_missing / n_total * 100), 1) if n_total > 0 else 0.0
    series = col.dropna().astype(str)

    if len(series) == 0:
        return {
            "type": "categorical",
            "variable": variable_name,
            "n_total": n_total,
            "n_missing": n_missing,
            "pct_missing": pct_missing,
            "n_unique": 0,
            "mode": "",
            "frequencies": [],
        }

    vc = series.value_counts()
    n_valid = len(series)
    frequencies = [
        {
            "category": str(k),
            "n": int(v),
            "pct": round(float(v / n_valid * 100), 1),
        }
        for k, v in vc.items()
    ]
    mode_val = str(series.mode().iloc[0]) if len(series.mode()) > 0 else ""

    return {
        "type": "categorical",
        "variable": variable_name,
        "n_total": n_total,
        "n_missing": n_missing,
        "pct_missing": pct_missing,
        "n_unique": int(series.nunique()),
        "mode": mode_val,
        "frequencies": frequencies,
    }


def compute_variable_stats_continuous(
    df: pd.DataFrame, variable_name: str
) -> Dict[str, Any]:
    if variable_name not in df.columns:
        raise ValueError(f"Column '{variable_name}' not found in dataset")

    col = df[variable_name]
    n_total = int(len(col))
    n_missing = int(col.isna().sum())
    pct_missing = round(float(n_missing / n_total * 100), 1) if n_total > 0 else 0.0

    try:
        series = pd.to_numeric(col, errors="coerce").dropna()
    except Exception:
        series = col.dropna()

    if len(series) == 0:
        return {
            "type": "continuous",
            "variable": variable_name,
            "n_total": n_total,
            "n_missing": n_missing,
            "pct_missing": pct_missing,
            "mean": None,
            "sd": None,
            "median": None,
            "q1": None,
            "q3": None,
            "min": None,
            "max": None,
            "skewness": None,
            "kurtosis": None,
            "histogram_bins": [],
        }

    q1 = float(series.quantile(0.25))
    q3 = float(series.quantile(0.75))
    skewness = float(series.skew()) if len(series) > 2 else 0.0
    kurt = float(series.kurtosis()) if len(series) > 3 else 0.0
    n_bins = min(15, max(10, int(np.sqrt(len(series)))))

    return {
        "type": "continuous",
        "variable": variable_name,
        "n_total": n_total,
        "n_missing": n_missing,
        "pct_missing": pct_missing,
        "mean": round(float(series.mean()), 3),
        "sd": round(float(series.std()), 3),
        "median": round(float(series.median()), 3),
        "q1": round(q1, 3),
        "q3": round(q3, 3),
        "min": round(float(series.min()), 3),
        "max": round(float(series.max()), 3),
        "skewness": round(skewness, 3),
        "kurtosis": round(kurt, 3),
        "histogram_bins": _get_histogram_bins(series, n_bins),
    }


# ---------------------------------------------------------------------------
# Table 1 builder
# ---------------------------------------------------------------------------

def build_table1(
    df: pd.DataFrame,
    variables: List[Dict[str, str]],
    summary_type: str,
) -> Dict[str, Any]:
    n_total = int(len(df))
    table: List[Dict[str, Any]] = []

    for var in variables:
        var_name = var["name"]
        var_type = var["type"]

        if var_name not in df.columns:
            continue

        if var_type == "categorical":
            col = df[var_name].dropna().astype(str)
            vc = col.value_counts()
            n_valid = len(col)
            rows = (
                [
                    {
                        "label": str(k),
                        "value": f"{int(v)} ({round(float(v) / n_valid * 100, 1)}%)",
                    }
                    for k, v in vc.items()
                ]
                if n_valid > 0
                else []
            )
            table.append(
                {
                    "variable": var_name,
                    "type": "categorical",
                    "summary_type": "n_pct",
                    "rows": rows,
                }
            )
        else:
            col_numeric = pd.to_numeric(df[var_name], errors="coerce").dropna()
            if len(col_numeric) == 0:
                continue

            skewness = float(col_numeric.skew()) if len(col_numeric) > 2 else 0.0
            use_median = (
                summary_type == "median_iqr"
                or (summary_type == "auto" and abs(skewness) > 1)
            )

            if use_median:
                q1 = round(float(col_numeric.quantile(0.25)), 1)
                q3 = round(float(col_numeric.quantile(0.75)), 1)
                med = round(float(col_numeric.median()), 1)
                value = f"{med} [{q1}\u2013{q3}]"
                label = f"{var_name}, median [IQR]"
                actual_type = "median_iqr"
            else:
                mean = round(float(col_numeric.mean()), 1)
                sd = round(float(col_numeric.std()), 1)
                value = f"{mean} \u00b1 {sd}"
                label = f"{var_name}, mean \u00b1 SD"
                actual_type = "mean_sd"

            table.append(
                {
                    "variable": var_name,
                    "type": "continuous",
                    "summary_type": actual_type,
                    "rows": [{"label": label, "value": value}],
                }
            )

    return {"n_total": n_total, "table": table}


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------

def generate_table1_docx(n_total: int, table: List[Dict[str, Any]]) -> bytes:
    from docx import Document  # type: ignore
    from docx.shared import Pt  # type: ignore

    doc = Document()
    doc.add_heading("Table 1. Baseline Characteristics", level=1)

    dtable = doc.add_table(rows=1, cols=2)
    dtable.style = "Table Grid"

    # Header
    hdr = dtable.rows[0].cells
    hdr[0].text = "Characteristic"
    hdr[1].text = f"Overall (N={n_total})"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    # Data rows
    for row_data in table:
        if row_data["type"] == "continuous":
            for r in row_data["rows"]:
                dr = dtable.add_row()
                dr.cells[0].text = r["label"]
                dr.cells[1].text = r["value"]
        else:
            # Bold variable label row
            lr = dtable.add_row()
            lp = lr.cells[0].paragraphs[0]
            lp.add_run(row_data["variable"]).bold = True
            lr.cells[1].text = ""
            # Indented category rows
            for r in row_data["rows"]:
                cr = dtable.add_row()
                cp = cr.cells[0].paragraphs[0]
                cp.paragraph_format.left_indent = Pt(12)
                cp.add_run(f"  {r['label']}")
                cr.cells[1].text = r["value"]

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
